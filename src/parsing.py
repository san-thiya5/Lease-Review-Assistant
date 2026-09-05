"""
Document text extraction module supporting PDF, DOCX, and TXT files.
Preserves paragraph breaks and clause headers while normalizing layout artifacts.
"""

import os
from pathlib import Path
import fitz  # PyMuPDF
import docx  # python-docx


class ParsingException(Exception):
    """Base exception for parsing failures."""
    pass


class UnsupportedFileTypeError(ParsingException):
    """Raised when a file with an unsupported extension is provided."""
    pass


class EmptyDocumentError(ParsingException):
    """Raised when the parsed document yields no readable text."""
    pass


class DocumentParsingError(ParsingException):
    """Raised when extraction fails due to file corruption or unreadable structure."""
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz).
    
    Layout edge cases handled:
    - Multi-column & reading order: Uses page.get_text("blocks") to preserve natural vertical
      reading order and avoid interleaving side-by-side columns.
    - Page headers and footers: Identifies and strips common recurring marginal page numbers
      (e.g., 'Page 1 of 5' or single digit lines near page margins).
    - Paragraph spacing: Inserts clean double-newlines between text blocks to preserve clause
      boundaries for downstream regex segmentation.
    """
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise DocumentParsingError(f"Failed to open PDF file {file_path}: {e}") from e

    extracted_blocks = []
    
    try:
        for page_num, page in enumerate(doc):
            # "blocks" returns tuples: (x0, y0, x1, y1, text, block_no, block_type)
            # block_type 0 is text
            blocks = page.get_text("blocks")
            for b in blocks:
                if b[6] == 0:  # text block
                    block_text = b[4].strip()
                    if not block_text:
                        continue
                    
                    # Filter out simple page number footers/headers (e.g., "Page 1", "- 2 -", single numbers)
                    lines = block_text.splitlines()
                    if len(lines) == 1 and (
                        lines[0].lower().startswith("page ") or 
                        lines[0].strip(" -_").isdigit()
                    ):
                        continue
                    
                    extracted_blocks.append(block_text)
    finally:
        doc.close()

    text = "\n\n".join(extracted_blocks).strip()
    return text


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX document using python-docx.
    
    Preserves:
    - Paragraph breaks: Joins non-empty paragraphs with double newlines.
    - Tables: Iterates over rows and cells, joining cell contents cleanly so
      lease terms presented in summary tables (e.g. rent tables) are retained.
    """
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise DocumentParsingError(f"Failed to open DOCX file {file_path}: {e}") from e

    elements = []

    # Extract regular paragraphs
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if p_text:
            elements.append(p_text)

    # Extract tables if present
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                elements.append(" | ".join(row_cells))

    text = "\n\n".join(elements).strip()
    return text


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a plain text or markdown file.
    Attempts UTF-8 first, falling back to latin-1 for legacy encodings.
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read().strip()
        except Exception as e:
            raise DocumentParsingError(f"Failed to read TXT file with fallback encoding: {e}") from e
    except Exception as e:
        raise DocumentParsingError(f"Failed to read TXT file {file_path}: {e}") from e


def extract_text(file_path: str) -> str:
    """
    Extract raw text from a lease document (.pdf, .docx, or .txt).
    
    Raises:
    - FileNotFoundError: If file does not exist.
    - UnsupportedFileTypeError: If file extension is not supported.
    - EmptyDocumentError: If document contains no readable text.
    - DocumentParsingError: If file is corrupted or unparseable.
    """
    path = Path(file_path)
    if not path.is_file():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(str(path))
    elif ext in [".docx", ".doc"]:
        if ext == ".doc":
            raise UnsupportedFileTypeError("Legacy .doc format is not supported. Please convert to .docx or .pdf.")
        raw_text = extract_text_from_docx(str(path))
    elif ext in [".txt", ".md"]:
        raw_text = extract_text_from_txt(str(path))
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file format '{ext}'. Supported formats are: .pdf, .docx, .txt"
        )

    # Normalize line endings
    raw_text = raw_text.replace("\r\n", "\n").replace("\r", "\n").strip()

    if not raw_text:
        raise EmptyDocumentError(f"Extracted document content from '{path.name}' is empty.")

    return raw_text
